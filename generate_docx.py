"""Generate DOCUMENTATION.docx from content."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, code, title=None):
    """Add a code block with monospace font."""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.style = 'No Spacing'
    for line in code.strip().split('\n'):
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(37, 99, 235)

def create_documentation():
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title
    title = doc.add_heading('Prompt Evaluation Pipeline', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)

    version = doc.add_paragraph('Version 0.1.0')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    desc = doc.add_paragraph('A Declarative YAML-Based Framework for Testing and Comparing LLM Prompts')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.runs[0].italic = True

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '   1.1 What is Prompt Engineering?',
        '   1.2 What is Prompt Evaluation?',
        '   1.3 Why Use This Framework?',
        '2. Quick Start',
        '3. Installation',
        '4. Architecture Overview',
        '5. CLI Reference',
        '6. Configuration Reference',
        '7. Metrics Reference',
        '8. Tutorial: Building a Q&A Chatbot Evaluation',
        '9. Sequence Diagrams',
        '10. Troubleshooting',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # Chapter 1: Introduction
    doc.add_heading('1. Introduction', level=1)

    doc.add_heading('1.1 What is Prompt Engineering?', level=2)
    doc.add_paragraph(
        'Prompt engineering is the practice of designing, refining, and optimizing the instructions (prompts) '
        'given to Large Language Models (LLMs) to achieve desired outputs. It involves:'
    )

    bullet_points = [
        ('Crafting clear instructions', 'Writing precise directions that guide the model\'s behavior'),
        ('Providing context', 'Including relevant background information to improve response quality'),
        ('Setting constraints', 'Defining boundaries for format, length, tone, and content'),
        ('Using examples', 'Demonstrating expected input-output patterns (few-shot learning)'),
        ('Iterative refinement', 'Testing and improving prompts based on observed results'),
    ]
    for title, desc in bullet_points:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        'Prompt engineering is crucial because the same model can produce vastly different results depending '
        'on how instructions are framed. A well-engineered prompt can be the difference between a helpful, '
        'accurate response and a vague or incorrect one.'
    )

    doc.add_heading('Example of Prompt Engineering Evolution', level=3)

    add_code_block(doc, 'Version 1 (Basic):\n"Answer the user\'s question."', 'Version 1:')
    add_code_block(doc,
        'Version 2 (With context):\n"You are a helpful customer support assistant. Answer the user\'s\n'
        'question clearly and concisely."', 'Version 2:')
    add_code_block(doc,
        'Version 3 (With constraints):\n"You are a helpful customer support assistant for TechCorp. Answer\n'
        'the user\'s question in 2-3 sentences. If you don\'t know the answer,\n'
        'say \'I\'ll need to check with our team.\' Never make up information."', 'Version 3:')

    doc.add_heading('1.2 What is Prompt Evaluation?', level=2)
    doc.add_paragraph(
        'Prompt evaluation is the systematic process of testing prompts against predefined criteria to measure '
        'their quality, consistency, and effectiveness. It answers the question: "How well does this prompt perform?"'
    )

    doc.add_paragraph('Prompt evaluation involves:')
    eval_points = [
        ('Defining test cases', 'Creating representative inputs with expected outcomes'),
        ('Running evaluations', 'Executing prompts against test cases and capturing outputs'),
        ('Measuring quality', 'Applying metrics to quantify how well outputs meet expectations'),
        ('Comparing versions', 'Testing multiple prompt variants to identify the best performer'),
        ('Tracking regression', 'Ensuring prompt changes don\'t degrade quality over time'),
    ]
    for title, desc in eval_points:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('Why is Prompt Evaluation Important?', level=3)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Problem'
    headers[1].text = 'Solution through Evaluation'
    set_cell_shading(headers[0], 'E5E7EB')
    set_cell_shading(headers[1], 'E5E7EB')

    data = [
        ('Inconsistent outputs', 'Measure consistency across multiple test cases'),
        ('Quality degradation after changes', 'Run regression tests before deployment'),
        ('Difficulty choosing between prompt versions', 'Compare metrics side-by-side'),
        ('No objective quality measurement', 'Apply quantifiable metrics (accuracy, relevance)'),
        ('Manual testing is time-consuming', 'Automate with declarative test configurations'),
    ]
    for i, (prob, sol) in enumerate(data):
        row = table.rows[i + 1].cells
        row[0].text = prob
        row[1].text = sol

    doc.add_heading('1.3 Why Use This Framework?', level=2)
    doc.add_paragraph('The Prompt Evaluation Pipeline provides:')

    features = [
        ('Declarative configuration', 'Define evaluations in YAML without writing code'),
        ('Multiple metrics', 'Built-in support for exact match, substring search, length validation, regex patterns, and LLM-based judging'),
        ('Batch evaluation', 'Run multiple test cases automatically'),
        ('Prompt comparison', 'Evaluate multiple prompt versions side-by-side'),
        ('Multiple export formats', 'Generate reports in JSON, CSV, Markdown, and HTML'),
        ('Prompt caching', 'Optimize API costs with Anthropic\'s prompt caching feature'),
        ('Extensible design', 'Easy to add custom metrics and exporters'),
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # Chapter 2: Quick Start
    doc.add_heading('2. Quick Start', level=1)
    doc.add_paragraph('Get started in 5 minutes:')

    doc.add_heading('Step 1: Install the package', level=2)
    add_code_block(doc, 'pip install -e .')

    doc.add_heading('Step 2: Set your API key', level=2)
    add_code_block(doc, 'export ANTHROPIC_API_KEY=your_key_here')

    doc.add_heading('Step 3: Create a prompt file', level=2)
    doc.add_paragraph('Create prompts/qa_assistant.yaml:')
    add_code_block(doc, '''name: qa_assistant_v1
version: "1.0"
description: "A helpful Q&A assistant"
system: |
  You are a helpful Q&A assistant. Answer questions accurately
  and concisely. If you don't know the answer, say "I don't know."
template: |
  Question: {question}''')

    doc.add_heading('Step 4: Create an evaluation config', level=2)
    doc.add_paragraph('Create configs/eval_qa.yaml:')
    add_code_block(doc, '''model: claude-sonnet-4-20250514
max_tokens: 256
temperature: 0.0
evaluation_threshold: 0.8

prompts:
  - prompts/qa_assistant.yaml

test_cases:
  - name: capital_france
    inputs:
      question: "What is the capital of France?"
    expected_contains:
      - "Paris"

  - name: unknown_question
    inputs:
      question: "What color is the invisible dragon?"
    expected_contains:
      - "don't know"

metrics:
  - type: contains
    case_sensitive: false''')

    doc.add_heading('Step 5: Run the evaluation', level=2)
    add_code_block(doc, 'prompt-eval run configs/eval_qa.yaml')

    doc.add_heading('Step 6: View results', level=2)
    doc.add_paragraph('Results are saved in results/qa/<timestamp>/ with JSON, CSV, Markdown, and HTML reports.')

    doc.add_page_break()

    # Chapter 3: Installation
    doc.add_heading('3. Installation', level=1)

    doc.add_heading('3.1 Requirements', level=2)
    doc.add_paragraph('• Python 3.10 or higher', style='List Bullet')
    doc.add_paragraph('• An Anthropic API key', style='List Bullet')

    doc.add_heading('3.2 Installation Steps', level=2)

    doc.add_heading('Option A: Install from source (recommended)', level=3)
    add_code_block(doc, '''# Clone the repository
git clone https://github.com/your-org/prompt-eval-pipeline.git
cd prompt-eval-pipeline

# Create a virtual environment
python -m venv venv
source venv/bin/activate  # On Windows: venv\\Scripts\\activate

# Install in editable mode
pip install -e .

# Install development dependencies (optional)
pip install -e ".[dev]"''')

    doc.add_heading('Option B: Install dependencies manually', level=3)
    add_code_block(doc, 'pip install anthropic>=0.40.0 pydantic>=2.0.0 pyyaml>=6.0 rich>=13.0.0 pandas>=2.0.0 jinja2>=3.0.0 click>=8.0.0 python-dotenv>=1.0.0 pypdf>=5.0.0')

    doc.add_heading('3.3 Configuration', level=2)
    doc.add_paragraph('Setting the API Key:')
    doc.add_paragraph('Option 1: Environment variable')
    add_code_block(doc, 'export ANTHROPIC_API_KEY=sk-ant-...')
    doc.add_paragraph('Option 2: Create a .env file in the project root')
    add_code_block(doc, 'ANTHROPIC_API_KEY=sk-ant-...')

    doc.add_heading('Verify installation', level=3)
    add_code_block(doc, 'prompt-eval --version\n# Output: prompt-eval, version 0.1.0')

    doc.add_page_break()

    # Chapter 4: Architecture
    doc.add_heading('4. Architecture Overview', level=1)
    doc.add_paragraph('The framework follows a modular architecture with clear separation of concerns:')

    # Architecture diagram as text
    arch_diagram = '''
┌─────────────────────────────────────────────────────────────────┐
│                         CLI Layer                               │
│                      (cli.py - Click)                           │
└─────────────────────────┬───────────────────────────────────────┘
                          │
┌─────────────────────────▼───────────────────────────────────────┐
│                    Configuration Layer                          │
│                 (loaders.py - YAML/JSON)                        │
└─────────────────────────┬───────────────────────────────────────┘
                          │
┌─────────────────────────▼───────────────────────────────────────┐
│                    Evaluation Engine                            │
│                (evaluator.py - Anthropic API)                   │
└───────────┬─────────────────────────────────┬───────────────────┘
            │                                 │
┌───────────▼───────────┐       ┌─────────────▼─────────────────┐
│     Metrics Layer     │       │       Export Layer            │
│     (metrics.py)      │       │     (exporters.py)            │
│ - Contains            │       │ - JSON                        │
│ - ExactMatch          │       │ - CSV                         │
│ - ResponseLength      │       │ - Markdown                    │
│ - RegexMatch          │       │ - HTML                        │
│ - LLMJudge            │       └───────────────────────────────┘
│ - Composite           │
└───────────────────────┘
'''
    p = doc.add_paragraph()
    run = p.add_run(arch_diagram)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    doc.add_heading('Component Responsibilities', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Component'
    headers[1].text = 'File'
    headers[2].text = 'Responsibility'
    for cell in headers:
        set_cell_shading(cell, 'E5E7EB')

    components = [
        ('CLI', 'cli.py', 'Command-line interface, orchestration'),
        ('Loaders', 'loaders.py', 'Parse YAML/JSON configurations'),
        ('Models', 'models.py', 'Data structures (Prompt, TestCase, EvalResult)'),
        ('Evaluator', 'evaluator.py', 'Execute prompts, call Anthropic API'),
        ('Metrics', 'metrics.py', 'Score outputs against expectations'),
        ('Exporters', 'exporters.py', 'Generate reports in multiple formats'),
    ]
    for i, (comp, file, resp) in enumerate(components):
        row = table.rows[i + 1].cells
        row[0].text = comp
        row[1].text = file
        row[2].text = resp

    doc.add_page_break()

    # Chapter 5: CLI Reference
    doc.add_heading('5. CLI Reference', level=1)

    doc.add_heading('5.1 Main Command', level=2)
    add_code_block(doc, 'prompt-eval [OPTIONS] COMMAND [ARGS]')
    doc.add_paragraph('Global Options:')
    doc.add_paragraph('• --version: Show version and exit', style='List Bullet')
    doc.add_paragraph('• --help: Show help message and exit', style='List Bullet')

    doc.add_heading('5.2 Run Command', level=2)
    doc.add_paragraph('Execute an evaluation from a configuration file.')
    add_code_block(doc, 'prompt-eval run [OPTIONS] CONFIG_FILE')

    doc.add_paragraph('Arguments:')
    doc.add_paragraph('• CONFIG_FILE: Path to the evaluation configuration file (YAML or JSON)', style='List Bullet')

    doc.add_paragraph('Options:')
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Option'
    headers[1].text = 'Short'
    headers[2].text = 'Description'
    headers[3].text = 'Default'
    for cell in headers:
        set_cell_shading(cell, 'E5E7EB')

    options = [
        ('--output', '-o', 'Output directory for results', 'Auto-generated'),
        ('--format', '-f', 'Output format: json, csv, markdown, html, all', 'all'),
        ('--model', '-m', 'Override model from config', 'Config value'),
        ('--verbose', '-v', 'Enable verbose output', 'True'),
        ('--quiet', '-q', 'Disable verbose output', 'False'),
    ]
    for i, (opt, short, desc, default) in enumerate(options):
        row = table.rows[i + 1].cells
        row[0].text = opt
        row[1].text = short
        row[2].text = desc
        row[3].text = default

    doc.add_heading('5.3 Examples', level=2)
    add_code_block(doc, '''# Basic run
prompt-eval run configs/eval_qa.yaml

# Specify output directory
prompt-eval run configs/eval_qa.yaml -o results/experiment1

# Only generate JSON output
prompt-eval run configs/eval_qa.yaml -f json

# Override model
prompt-eval run configs/eval_qa.yaml -m claude-opus-4-5-20251101

# Quiet mode (minimal output)
prompt-eval run configs/eval_qa.yaml -q''')

    doc.add_page_break()

    # Chapter 6: Configuration Reference
    doc.add_heading('6. Configuration Reference', level=1)

    doc.add_heading('6.1 Evaluation Configuration', level=2)
    doc.add_paragraph('The main configuration file defines the evaluation parameters.')
    add_code_block(doc, '''# Model Configuration
model: claude-sonnet-4-20250514    # Required: Anthropic model ID
max_tokens: 1024                    # Optional: Max response tokens
temperature: 0.0                    # Optional: Sampling temperature
evaluation_threshold: 0.8           # Required: Pass/fail threshold (0.0-1.0)

# Prompts to evaluate
prompts:
  - prompts/assistant_v1.yaml
  - prompts/assistant_v2.yaml

# Test cases
test_cases:
  - name: test_1
    inputs:
      question: "What is 2+2?"
    expected_contains:
      - "4"

# Metrics configuration
metrics:
  - type: contains
    case_sensitive: false''')

    doc.add_heading('6.2 Prompt Configuration', level=2)
    add_code_block(doc, '''# Basic Information
name: qa_assistant_v1              # Required: Unique identifier
version: "1.0"                      # Optional: Version string
description: "Q&A assistant"        # Optional: Description

# Prompt Content
system: |                           # Optional: System prompt
  You are a helpful assistant.

template: |                         # Required: User prompt template
  Question: {question}

# Advanced Prompt Engineering
rules:                              # Optional: Behavioral rules
  - Always be polite
  - Never make up information

skills:                             # Optional: Skills/capabilities
  - Answering factual questions

thinking_process: |                 # Optional: Chain-of-thought
  1. Understand the question
  2. Formulate a clear answer

examples:                           # Optional: Few-shot examples
  - input: "What is the sun?"
    output: "The sun is a star."

# File Caching
cached_files:                       # Optional: Files to cache
  - reference/knowledge_base.md''')

    doc.add_heading('6.3 Configuration Parameters Summary', level=2)
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Parameter'
    headers[1].text = 'Required'
    headers[2].text = 'Type'
    headers[3].text = 'Description'
    for cell in headers:
        set_cell_shading(cell, 'E5E7EB')

    params = [
        ('model', 'Yes', 'string', 'Anthropic model ID'),
        ('max_tokens', 'No', 'int', 'Max response tokens (default: 1024)'),
        ('temperature', 'No', 'float', 'Sampling temperature (default: 0.0)'),
        ('evaluation_threshold', 'Yes', 'float', 'Pass/fail threshold 0.0-1.0'),
        ('prompts', 'Yes', 'list', 'Prompt file paths or inline definitions'),
        ('test_cases', 'Yes', 'list', 'Test case definitions'),
        ('metrics', 'No', 'list', 'Metric configurations'),
    ]
    for i, (param, req, typ, desc) in enumerate(params):
        row = table.rows[i + 1].cells
        row[0].text = param
        row[1].text = req
        row[2].text = typ
        row[3].text = desc

    doc.add_page_break()

    # Chapter 7: Metrics Reference
    doc.add_heading('7. Metrics Reference', level=1)
    doc.add_paragraph('Metrics measure how well the LLM output matches expectations.')

    doc.add_heading('7.1 Contains Metric', level=2)
    doc.add_paragraph('Checks if the output contains required substrings and excludes forbidden ones.')
    add_code_block(doc, '''metrics:
  - type: contains
    case_sensitive: false    # Optional (default: false)''')
    doc.add_paragraph('Scoring: Score = min(matched_required/total_required, 1 - violations/total_forbidden)')

    doc.add_heading('7.2 Exact Match Metric', level=2)
    doc.add_paragraph('Checks if the output exactly matches the expected value.')
    add_code_block(doc, '''metrics:
  - type: exact_match
    case_sensitive: false    # Optional (default: false)
    strip: true              # Optional (default: true)''')
    doc.add_paragraph('Scoring: Binary (1.0 if match, 0.0 otherwise)')

    doc.add_heading('7.3 Response Length Metric', level=2)
    doc.add_paragraph('Validates the response length in characters or words.')
    add_code_block(doc, '''metrics:
  - type: response_length
    min_chars: 100           # Optional
    max_chars: 500           # Optional
    min_words: 20            # Optional
    max_words: 100           # Optional''')

    doc.add_heading('7.4 Regex Match Metric', level=2)
    doc.add_paragraph('Checks if the output matches a regular expression pattern.')
    add_code_block(doc, '''metrics:
  - type: regex_match
    pattern: "\\\\d{4}-\\\\d{2}-\\\\d{2}"    # Date pattern YYYY-MM-DD''')

    doc.add_heading('7.5 LLM Judge Metric', level=2)
    doc.add_paragraph('Uses another LLM to evaluate the quality of the output.')
    add_code_block(doc, '''metrics:
  - type: llm_judge
    criteria: |
      Evaluate if the response is:
      1. Accurate and factual
      2. Clear and well-organized
    model: claude-sonnet-4-20250514''')

    doc.add_heading('7.6 Metrics Summary', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Metric'
    headers[1].text = 'Use Case'
    headers[2].text = 'Scoring'
    for cell in headers:
        set_cell_shading(cell, 'E5E7EB')

    metrics = [
        ('contains', 'Check for required/forbidden content', 'Ratio of matches'),
        ('exact_match', 'Precise output matching', 'Binary (0 or 1)'),
        ('response_length', 'Enforce length constraints', 'Binary (0 or 1)'),
        ('regex_match', 'Pattern validation', 'Binary (0 or 1)'),
        ('llm_judge', 'Subjective quality assessment', 'Continuous (0-1)'),
        ('composite', 'Combined evaluation', 'Weighted average'),
    ]
    for i, (metric, use, score) in enumerate(metrics):
        row = table.rows[i + 1].cells
        row[0].text = metric
        row[1].text = use
        row[2].text = score

    doc.add_page_break()

    # Chapter 8: Tutorial
    doc.add_heading('8. Tutorial: Building a Q&A Chatbot Evaluation', level=1)

    doc.add_heading('8.1 Scenario', level=2)
    doc.add_paragraph("You're building a general-purpose Q&A chatbot that should:")
    doc.add_paragraph('• Answer factual questions accurately', style='List Bullet')
    doc.add_paragraph('• Admit when it doesn\'t know something', style='List Bullet')
    doc.add_paragraph('• Provide concise but complete answers', style='List Bullet')
    doc.add_paragraph('• Avoid making up information', style='List Bullet')

    doc.add_heading('8.2 Project Structure', level=2)
    add_code_block(doc, '''my-qa-chatbot/
├── prompts/
│   ├── qa_v1.yaml
│   └── qa_v2.yaml
├── configs/
│   └── eval_qa.yaml
└── results/''')

    doc.add_heading('8.3 Prompt Version 1 (Basic)', level=2)
    add_code_block(doc, '''name: qa_chatbot_v1
version: "1.0"
description: "Basic Q&A chatbot - first iteration"

system: |
  You are a helpful Q&A assistant. Answer user questions accurately.

template: |
  User Question: {question}''')

    doc.add_heading('8.4 Prompt Version 2 (Improved)', level=2)
    add_code_block(doc, '''name: qa_chatbot_v2
version: "2.0"
description: "Improved Q&A chatbot with better instructions"

system: |
  You are a knowledgeable and helpful Q&A assistant.

rules:
  - Provide accurate, factual information
  - Keep answers concise (2-4 sentences for simple questions)
  - If you don't know, say "I don't have information about that"
  - Never make up facts or statistics

thinking_process: |
  1. Understand what the user is asking
  2. Determine if I have reliable information
  3. Formulate a clear response or admit ignorance

template: |
  Please answer the following question:
  {question}''')

    doc.add_heading('8.5 Test Cases', level=2)
    add_code_block(doc, '''model: claude-sonnet-4-20250514
max_tokens: 512
temperature: 0.0
evaluation_threshold: 0.8

prompts:
  - prompts/qa_v1.yaml
  - prompts/qa_v2.yaml

test_cases:
  # Factual questions
  - name: capital_france
    inputs:
      question: "What is the capital of France?"
    expected_contains:
      - "Paris"
    tags: [geography, factual]

  - name: water_formula
    inputs:
      question: "What is the chemical formula for water?"
    expected_contains:
      - "H2O"
    tags: [science, factual]

  # Unknown questions
  - name: unknown_personal
    inputs:
      question: "What did I have for breakfast yesterday?"
    expected_contains:
      - "don't"
    expected_not_contains:
      - "you had"
    tags: [unknown, boundary]

metrics:
  - type: contains
    case_sensitive: false
  - type: response_length
    min_words: 5
    max_words: 200''')

    doc.add_heading('8.6 Run and Analyze', level=2)
    add_code_block(doc, 'prompt-eval run configs/eval_qa.yaml')
    doc.add_paragraph('Results will show which prompt version performs better across all test cases.')

    doc.add_page_break()

    # Chapter 9: Sequence Diagrams
    doc.add_heading('9. Sequence Diagrams', level=1)

    doc.add_heading('9.1 Single Prompt Evaluation Flow', level=2)
    seq1 = '''
┌─────┐          ┌─────┐          ┌────────┐          ┌─────────┐          ┌─────────┐
│User │          │ CLI │          │Loader  │          │Evaluator│          │Anthropic│
└──┬──┘          └──┬──┘          └───┬────┘          └────┬────┘          └────┬────┘
   │                │                 │                    │                    │
   │ prompt-eval    │                 │                    │                    │
   │ run config.yaml│                 │                    │                    │
   │───────────────>│                 │                    │                    │
   │                │                 │                    │                    │
   │                │ load_eval_config│                    │                    │
   │                │────────────────>│                    │                    │
   │                │                 │                    │                    │
   │                │     config      │                    │                    │
   │                │<────────────────│                    │                    │
   │                │                 │                    │                    │
   │                │     evaluate(prompt, test_cases)     │                    │
   │                │─────────────────────────────────────>│                    │
   │                │                 │                    │                    │
   │                │                 │                    │  ┌────────────────┐│
   │                │                 │                    │  │ For each       ││
   │                │                 │                    │  │ test case      ││
   │                │                 │                    │  └───────┬────────┘│
   │                │                 │                    │          │         │
   │                │                 │                    │ messages.create    │
   │                │                 │                    │─────────────────────>
   │                │                 │                    │          │         │
   │                │                 │                    │     response       │
   │                │                 │                    │<─────────────────────
   │                │                 │                    │          │         │
   │                │                 │                    │ calculate metrics  │
   │                │                 │                    │──────────┤         │
   │                │                 │                    │          │         │
   │                │      EvalRun (results)               │                    │
   │                │<─────────────────────────────────────│                    │
   │                │                 │                    │                    │
   │  Results saved │                 │                    │                    │
   │<───────────────│                 │                    │                    │
'''
    p = doc.add_paragraph()
    run = p.add_run(seq1)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)

    doc.add_heading('9.2 Multi-Prompt Comparison Flow', level=2)
    seq2 = '''
┌─────┐          ┌─────┐          ┌─────────┐          ┌────────┐
│User │          │ CLI │          │Evaluator│          │Exporter│
└──┬──┘          └──┬──┘          └────┬────┘          └───┬────┘
   │                │                  │                   │
   │ run config     │                  │                   │
   │ (2 prompts)    │                  │                   │
   │───────────────>│                  │                   │
   │                │                  │                   │
   │                │ evaluate(prompt1)│                   │
   │                │─────────────────>│                   │
   │                │                  │                   │
   │                │     EvalRun 1    │                   │
   │                │<─────────────────│                   │
   │                │                  │                   │
   │                │ evaluate(prompt2)│                   │
   │                │─────────────────>│                   │
   │                │                  │                   │
   │                │     EvalRun 2    │                   │
   │                │<─────────────────│                   │
   │                │                  │                   │
   │                │ export_all(run1) │                   │
   │                │─────────────────────────────────────>│
   │                │                  │                   │
   │                │ export_all(run2) │                   │
   │                │─────────────────────────────────────>│
   │                │                  │                   │
   │                │ compare_runs([run1, run2])           │
   │                │─────────────────────────────────────>│
   │                │                  │                   │
   │                │        comparison.md                 │
   │                │<─────────────────────────────────────│
   │                │                  │                   │
   │ Results +      │                  │                   │
   │ comparison     │                  │                   │
   │<───────────────│                  │                   │
'''
    p = doc.add_paragraph()
    run = p.add_run(seq2)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)

    doc.add_heading('9.3 Metrics Calculation Flow', level=2)
    seq3 = '''
┌─────────┐          ┌───────────┐          ┌────────┐          ┌──────────┐
│Evaluator│          │ Contains  │          │LLMJudge│          │Anthropic │
└────┬────┘          └─────┬─────┘          └───┬────┘          └─────┬────┘
     │                     │                    │                     │
     │ score(output,       │                    │                     │
     │       test_case)    │                    │                     │
     │────────────────────>│                    │                     │
     │                     │                    │                     │
     │                     │ Check expected_    │                     │
     │                     │ contains           │                     │
     │                     │──────────┐         │                     │
     │                     │          │         │                     │
     │                     │<─────────┘         │                     │
     │                     │                    │                     │
     │   contains_score    │                    │                     │
     │<────────────────────│                    │                     │
     │                     │                    │                     │
     │ score(output, test_case)                 │                     │
     │─────────────────────────────────────────>│                     │
     │                     │                    │                     │
     │                     │                    │ Build judge prompt  │
     │                     │                    │──────────┐          │
     │                     │                    │          │          │
     │                     │                    │<─────────┘          │
     │                     │                    │                     │
     │                     │                    │ messages.create     │
     │                     │                    │────────────────────>│
     │                     │                    │                     │
     │                     │                    │    score (0.0-1.0)  │
     │                     │                    │<────────────────────│
     │                     │                    │                     │
     │            llm_judge_score               │                     │
     │<────────────────────────────────────────│                     │
     │                     │                    │                     │
     │ Aggregate scores    │                    │                     │
     │ passed = all >= threshold?              │                     │
'''
    p = doc.add_paragraph()
    run = p.add_run(seq3)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)

    doc.add_page_break()

    # Chapter 10: Troubleshooting
    doc.add_heading('10. Troubleshooting', level=1)

    doc.add_heading('10.1 Common Errors', level=2)

    doc.add_heading('Error: "ANTHROPIC_API_KEY not set"', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    p.add_run('The API key is not configured.')
    p = doc.add_paragraph()
    run = p.add_run('Solution:')
    run.bold = True
    add_code_block(doc, '''# Option 1: Set environment variable
export ANTHROPIC_API_KEY=sk-ant-...

# Option 2: Create .env file
echo "ANTHROPIC_API_KEY=sk-ant-..." > .env''')

    doc.add_heading('Error: "No prompts found in config"', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    p.add_run("The configuration file doesn't specify any prompts.")
    p = doc.add_paragraph()
    run = p.add_run('Solution: ')
    run.bold = True
    p.add_run('Add a prompts section to your config:')
    add_code_block(doc, '''prompts:
  - prompts/my_prompt.yaml''')

    doc.add_heading('Error: "evaluation_threshold is required"', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    p.add_run('Missing required threshold parameter.')
    p = doc.add_paragraph()
    run = p.add_run('Solution: ')
    run.bold = True
    p.add_run('Add the threshold to your config:')
    add_code_block(doc, 'evaluation_threshold: 0.8')

    doc.add_heading('10.2 Performance Issues', level=2)

    doc.add_heading('Slow Evaluation Speed', level=3)
    doc.add_paragraph('Causes and solutions:')
    doc.add_paragraph('1. Large number of test cases: Run in batches or use a faster model', style='List Number')
    doc.add_paragraph('2. High max_tokens: Reduce if responses are shorter', style='List Number')
    doc.add_paragraph('3. No prompt caching: Use cached_files for large reference documents', style='List Number')

    doc.add_heading('High API Costs', level=3)
    doc.add_paragraph('Solutions:')
    doc.add_paragraph('1. Use prompt caching for repeated content', style='List Number')
    doc.add_paragraph('2. Use a cheaper model for development', style='List Number')
    doc.add_paragraph('3. Reduce test cases during development', style='List Number')

    doc.add_heading('10.3 Configuration Debugging', level=2)
    doc.add_paragraph('Validate your configuration:')
    add_code_block(doc, '''# Check YAML syntax
python -c "import yaml; yaml.safe_load(open('config.yaml'))"

# Run with verbose output
prompt-eval run config.yaml -v''')

    doc.add_paragraph('Common YAML mistakes:')
    add_code_block(doc, '''# Wrong: Missing quotes around special characters
expected_contains:
  - What's the answer?    # Error: apostrophe

# Correct: Quote strings with special characters
expected_contains:
  - "What's the answer?"''')

    doc.add_page_break()

    # Appendix
    doc.add_heading('Appendix: Reference Tables', level=1)

    doc.add_heading('A. Supported Models', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Model ID'
    headers[1].text = 'Description'
    for cell in headers:
        set_cell_shading(cell, 'E5E7EB')

    models = [
        ('claude-opus-4-5-20251101', 'Most capable, best for complex tasks'),
        ('claude-sonnet-4-20250514', 'Balanced performance and cost'),
        ('claude-haiku-3', 'Fastest, best for simple tasks'),
    ]
    for i, (model, desc) in enumerate(models):
        row = table.rows[i + 1].cells
        row[0].text = model
        row[1].text = desc

    doc.add_heading('B. File Format Support', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Format'
    headers[1].text = 'Extension'
    headers[2].text = 'Use Case'
    for cell in headers:
        set_cell_shading(cell, 'E5E7EB')

    formats = [
        ('YAML', '.yaml, .yml', 'Configuration files'),
        ('JSON', '.json', 'Configuration files'),
        ('Markdown', '.md', 'Reference documents'),
        ('Text', '.txt', 'Reference documents'),
        ('PDF', '.pdf', 'Reference documents'),
    ]
    for i, (fmt, ext, use) in enumerate(formats):
        row = table.rows[i + 1].cells
        row[0].text = fmt
        row[1].text = ext
        row[2].text = use

    doc.add_heading('C. Output Formats', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Format'
    headers[1].text = 'Extension'
    headers[2].text = 'Best For'
    for cell in headers:
        set_cell_shading(cell, 'E5E7EB')

    outputs = [
        ('JSON', '.json', 'Programmatic processing'),
        ('CSV', '.csv', 'Spreadsheet analysis'),
        ('Markdown', '.md', 'Documentation'),
        ('HTML', '.html', 'Visual review'),
    ]
    for i, (fmt, ext, use) in enumerate(outputs):
        row = table.rows[i + 1].cells
        row[0].text = fmt
        row[1].text = ext
        row[2].text = use

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Generated for Prompt Evaluation Pipeline v0.1.0')
    run.italic = True

    return doc

if __name__ == '__main__':
    doc = create_documentation()
    doc.save('DOCUMENTATION.docx')
    print('DOCUMENTATION.docx created successfully!')
